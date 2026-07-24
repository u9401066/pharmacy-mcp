"""Safe text extraction for configured local pharmaceutical files."""

from __future__ import annotations

import csv
import hashlib
import shutil

# Required only for optional legacy .doc extraction; never invoked through a shell.
import subprocess  # nosec B404
from dataclasses import dataclass
from pathlib import Path
from typing import Any

SUPPORTED_EXTENSIONS = frozenset(
    {".csv", ".doc", ".docx", ".md", ".pdf", ".txt", ".xls", ".xlsx"}
)
MAX_READ_CHARS = 50_000


@dataclass(frozen=True)
class DocumentMatch:
    document_id: str
    path: str
    relative_path: str
    title: str
    extension: str
    snippet: str
    line_start: int
    line_end: int
    char_start: int
    char_end: int
    text_sha256: str


@dataclass(frozen=True)
class DocumentRead:
    """One bounded, exact text span addressable without a caller-supplied path."""

    document_id: str
    path: str
    relative_path: str
    title: str
    extension: str
    content: str
    line_start: int
    line_end: int
    char_start: int
    char_end: int
    total_chars: int
    truncated: bool
    text_sha256: str


class DocumentStore:
    """Search files under administrator-configured roots only."""

    def __init__(
        self,
        roots: tuple[Path, ...],
        *,
        max_bytes: int,
        max_files: int,
    ) -> None:
        self.roots = tuple(root.resolve() for root in roots)
        self.max_bytes = max_bytes
        self.max_files = max_files

    def search(
        self,
        query: str,
        limit: int,
    ) -> tuple[list[DocumentMatch], list[str], int]:
        """Extract configured files and return bounded matching snippets."""

        matches: list[DocumentMatch] = []
        warnings: list[str] = []
        scanned = 0
        for path in self._files():
            if scanned >= self.max_files or len(matches) >= limit:
                break
            scanned += 1
            try:
                text = self.read(path)
            except (OSError, ValueError, RuntimeError) as exc:
                warnings.append(f"{path.name}: {exc}")
                continue
            span = _matching_span(text, query)
            if span is None:
                continue
            start, end = span
            line_start, line_end = _line_range(text, start, end)
            document_id, relative_path = self._document_identity(path)
            matches.append(
                DocumentMatch(
                    document_id=document_id,
                    path=str(path),
                    relative_path=relative_path,
                    title=path.stem,
                    extension=path.suffix.lower(),
                    snippet=text[start:end],
                    line_start=line_start,
                    line_end=line_end,
                    char_start=start,
                    char_end=end,
                    text_sha256=_text_sha256(text),
                )
            )
        return matches, warnings, scanned

    def read_by_id(
        self,
        document_id: str,
        *,
        offset: int = 0,
        max_chars: int = 10_000,
    ) -> DocumentRead:
        """Return one exact, bounded span identified by a search result ID."""

        if not document_id.startswith("doc-"):
            raise ValueError("invalid document ID")
        if offset < 0:
            raise ValueError("offset must be zero or greater")
        if not 1 <= max_chars <= MAX_READ_CHARS:
            raise ValueError(f"max_chars must be between 1 and {MAX_READ_CHARS}")

        path = self._path_for_id(document_id)
        if path is None:
            raise ValueError("document ID was not found in configured roots")
        text = self.read(path)
        if offset > len(text):
            raise ValueError("offset exceeds extracted document length")
        end = min(len(text), offset + max_chars)
        line_start, line_end = _line_range(text, offset, end)
        resolved_id, relative_path = self._document_identity(path)
        return DocumentRead(
            document_id=resolved_id,
            path=str(path),
            relative_path=relative_path,
            title=path.stem,
            extension=path.suffix.lower(),
            content=text[offset:end],
            line_start=line_start,
            line_end=line_end,
            char_start=offset,
            char_end=end,
            total_chars=len(text),
            truncated=end < len(text),
            text_sha256=_text_sha256(text),
        )

    def read(self, path: Path) -> str:
        """Extract text after enforcing root, file type, symlink, and size policy."""

        resolved = path.resolve()
        if path.is_symlink() or not self._inside_root(resolved):
            raise ValueError("path is outside an allowed root or is a symlink")
        if resolved.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError("unsupported file extension")
        if resolved.stat().st_size > self.max_bytes:
            raise ValueError(f"file exceeds {self.max_bytes} byte limit")

        extension = resolved.suffix.lower()
        if extension in {".md", ".txt"}:
            return resolved.read_text(encoding="utf-8", errors="replace")
        if extension == ".csv":
            return _read_csv(resolved)
        if extension == ".pdf":
            return _read_pdf(resolved)
        if extension == ".docx":
            return _read_docx(resolved)
        if extension == ".xlsx":
            return _read_xlsx(resolved)
        if extension == ".xls":
            return _read_xls(resolved)
        return _read_legacy_doc(resolved)

    def _files(self) -> list[Path]:
        files: set[Path] = set()
        for root in self.roots:
            if not root.is_dir():
                continue
            files.update(
                path
                for path in root.rglob("*")
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
            )
        return sorted(files)

    def _inside_root(self, path: Path) -> bool:
        return any(path.is_relative_to(root) for root in self.roots)

    def _path_for_id(self, document_id: str) -> Path | None:
        for path in self._files():
            try:
                candidate, _ = self._document_identity(path)
            except ValueError:
                continue
            if candidate == document_id:
                return path
        return None

    def _document_identity(self, path: Path) -> tuple[str, str]:
        resolved = path.resolve()
        if path.is_symlink():
            raise ValueError("symlinks cannot be assigned a document ID")
        for root_index, root in enumerate(self.roots):
            if resolved.is_relative_to(root):
                relative_path = resolved.relative_to(root).as_posix()
                identity = f"{root_index}:{relative_path}".encode()
                digest = hashlib.sha256(identity).hexdigest()[:24]
                return f"doc-{digest}", relative_path
        raise ValueError("path is outside configured roots")


def _read_csv(path: Path) -> str:
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as file:
        return "\n".join(" | ".join(row) for row in csv.reader(file))


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_rows = [
        " | ".join(cell.text for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    return "\n".join([*paragraphs, *table_rows])


def _read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        rows = []
        for sheet in workbook.worksheets:
            rows.append(f"[{sheet.title}]")
            rows.extend(
                _stringify_row(row) for row in sheet.iter_rows(values_only=True)
            )
        return "\n".join(rows)
    finally:
        workbook.close()


def _read_xls(path: Path) -> str:
    import xlrd

    workbook = xlrd.open_workbook(str(path), on_demand=True)
    try:
        rows = []
        for sheet in workbook.sheets():
            rows.append(f"[{sheet.name}]")
            rows.extend(
                _stringify_row(sheet.row_values(index)) for index in range(sheet.nrows)
            )
        return "\n".join(rows)
    finally:
        workbook.release_resources()


def _read_legacy_doc(path: Path) -> str:
    executable = shutil.which("antiword")
    if executable is None:
        raise RuntimeError("legacy .doc requires the antiword executable")
    try:
        # The executable is resolved to an absolute path and the validated file
        # path is passed as one argument with shell=False.
        result = subprocess.run(  # nosec B603
            [executable, str(path)],
            check=True,
            capture_output=True,
            timeout=20,
        )
    except subprocess.SubprocessError as exc:
        raise RuntimeError(f"antiword extraction failed: {exc}") from exc
    return result.stdout.decode("utf-8", errors="replace")


def _stringify_row(row: Any) -> str:
    return " | ".join("" if value is None else str(value) for value in row)


def _matching_span(
    text: str,
    query: str,
    size: int = 500,
) -> tuple[int, int] | None:
    haystack = text.casefold()
    needle = query.casefold().strip()
    if not needle:
        return None
    index = haystack.find(needle)
    if index < 0:
        return None
    start = max(0, index - size // 3)
    end = min(len(text), start + size)
    return start, end


def _line_range(text: str, start: int, end: int) -> tuple[int, int]:
    line_start = text.count("\n", 0, start) + 1
    if end <= start:
        return line_start, line_start
    return line_start, text.count("\n", 0, end - 1) + 1


def _text_sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()
